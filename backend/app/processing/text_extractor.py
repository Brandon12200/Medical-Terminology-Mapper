"""
Text extraction module for various document formats
"""

import os
import logging
from typing import Optional, Dict, Any, Tuple
from pathlib import Path
import PyPDF2
from docx import Document
import tika
from tika import parser
import re

from app.utils.logger import setup_logger

# Initialize Tika server path to avoid issues
tika.initVM()

logger = setup_logger(__name__)


class TextExtractor:
    """Extract text content from various document formats"""
    
    def __init__(self):
        """Initialize the text extractor"""
        self.extraction_methods = {
            'pdf': self._extract_pdf,
            'docx': self._extract_docx,
            'txt': self._extract_txt,
            'rtf': self._extract_rtf,
            'hl7': self._extract_hl7
        }
    
    def extract_text(self, file_path: str, document_type: str) -> Tuple[Optional[str], str, Optional[Dict[str, Any]]]:
        """
        Extract text from a document file
        
        Args:
            file_path: Path to the document file
            document_type: Type of document (pdf, docx, txt, rtf, hl7)
            
        Returns:
            Tuple of (extracted_text, extraction_method, metadata)
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                logger.error(f"File not found: {file_path}")
                return None, "error", {"error": "File not found"}
            
            # Get the appropriate extraction method
            extraction_method = self.extraction_methods.get(document_type.lower())
            if not extraction_method:
                logger.error(f"Unsupported document type: {document_type}")
                return None, "error", {"error": f"Unsupported document type: {document_type}"}
            
            # Extract text
            text, method, metadata = extraction_method(file_path)
            
            # Clean up the extracted text
            if text:
                text = self._clean_text(text)
            
            return text, method, metadata
            
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return None, "error", {"error": str(e)}
    
    def _extract_pdf(self, file_path: Path) -> Tuple[Optional[str], str, Dict[str, Any]]:
        """Extract text from PDF files"""
        try:
            text_parts = []
            metadata = {"page_count": 0}
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata["page_count"] = len(pdf_reader.pages)
                
                # Extract metadata
                if pdf_reader.metadata:
                    metadata["pdf_metadata"] = {
                        "title": pdf_reader.metadata.get('/Title', ''),
                        "author": pdf_reader.metadata.get('/Author', ''),
                        "subject": pdf_reader.metadata.get('/Subject', ''),
                        "creator": pdf_reader.metadata.get('/Creator', '')
                    }
                
                # Extract text from each page
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num + 1}: {e}")
            
            text = '\n\n'.join(text_parts)
            
            # If PyPDF2 extraction fails or returns very little text, try Tika
            if not text or len(text.strip()) < 100:
                logger.info("PyPDF2 extraction yielded little text, trying Tika")
                return self._extract_with_tika(file_path, metadata)
            
            return text, "PyPDF2", metadata
            
        except Exception as e:
            logger.warning(f"PyPDF2 extraction failed, trying Tika: {e}")
            return self._extract_with_tika(file_path, {"error": str(e)})
    
    def _extract_docx(self, file_path: Path) -> Tuple[Optional[str], str, Dict[str, Any]]:
        """Extract text from DOCX files"""
        try:
            doc = Document(file_path)
            text_parts = []
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }
            
            # Extract core properties
            core_props = doc.core_properties
            if core_props:
                metadata["docx_metadata"] = {
                    "title": core_props.title or '',
                    "author": core_props.author or '',
                    "subject": core_props.subject or '',
                    "created": str(core_props.created) if core_props.created else ''
                }
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            text = '\n\n'.join(text_parts)
            return text, "python-docx", metadata
            
        except Exception as e:
            logger.warning(f"python-docx extraction failed, trying Tika: {e}")
            return self._extract_with_tika(file_path, {"error": str(e)})
    
    def _extract_txt(self, file_path: Path) -> Tuple[Optional[str], str, Dict[str, Any]]:
        """Extract text from TXT files"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                        metadata = {
                            "encoding": encoding,
                            "file_size": file_path.stat().st_size,
                            "line_count": len(text.splitlines())
                        }
                        return text, "direct_read", metadata
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, try Tika
            logger.warning("All encoding attempts failed, trying Tika")
            return self._extract_with_tika(file_path, {"encoding_error": True})
            
        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            return None, "error", {"error": str(e)}
    
    def _extract_rtf(self, file_path: Path) -> Tuple[Optional[str], str, Dict[str, Any]]:
        """Extract text from RTF files using Tika"""
        # RTF extraction is complex, use Tika directly
        return self._extract_with_tika(file_path, {"format": "rtf"})
    
    def _extract_hl7(self, file_path: Path) -> Tuple[Optional[str], str, Dict[str, Any]]:
        """Extract and parse HL7 messages"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Basic HL7 parsing
            segments = content.split('\r')
            if not segments:
                segments = content.split('\n')
            
            metadata = {
                "segment_count": len(segments),
                "message_type": None,
                "message_id": None
            }
            
            # Parse message header
            for segment in segments:
                if segment.startswith('MSH'):
                    fields = segment.split('|')
                    if len(fields) > 8:
                        metadata["message_type"] = fields[8]
                    if len(fields) > 9:
                        metadata["message_id"] = fields[9]
                    break
            
            # For HL7, we'll return the original content as "text"
            # Real applications would parse this into structured data
            return content, "hl7_parser", metadata
            
        except Exception as e:
            logger.warning(f"HL7 extraction failed, trying Tika: {e}")
            return self._extract_with_tika(file_path, {"error": str(e)})
    
    def _extract_with_tika(self, file_path: Path, initial_metadata: Optional[Dict[str, Any]] = None) -> Tuple[Optional[str], str, Dict[str, Any]]:
        """Extract text using Apache Tika as a fallback"""
        try:
            # Parse the document with Tika
            parsed = parser.from_file(str(file_path))
            
            text = parsed.get('content', '')
            metadata = initial_metadata or {}
            
            # Add Tika metadata
            if parsed.get('metadata'):
                metadata['tika_metadata'] = parsed['metadata']
            
            return text, "Apache Tika", metadata
            
        except Exception as e:
            logger.error(f"Tika extraction failed: {e}")
            return None, "error", {"error": str(e), "initial_metadata": initial_metadata}
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove non-printable characters (except newlines and tabs)
        text = ''.join(char for char in text if char.isprintable() or char in '\n\t')
        
        # Normalize line breaks
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Remove leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def get_text_preview(self, text: str, max_length: int = 500) -> str:
        """Get a preview of the extracted text"""
        if not text:
            return ""
        
        if len(text) <= max_length:
            return text
        
        # Find a good break point (end of sentence or word)
        preview = text[:max_length]
        
        # Try to break at sentence
        last_period = preview.rfind('.')
        if last_period > max_length * 0.8:
            return preview[:last_period + 1]
        
        # Break at word
        last_space = preview.rfind(' ')
        if last_space > max_length * 0.8:
            return preview[:last_space] + "..."
        
        return preview + "..."